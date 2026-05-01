from dataclasses import dataclass
from pathlib import Path

import docx
import fitz


@dataclass
class DocumentPage:
    """
    Represents text extracted from a single logical page.

    For .txt, .md, and .docx files, page_number is usually 1.
    For PDFs, page_number matches the PDF page.
    """

    page_number: int
    text: str


class DocumentLoader:
    """
    Extracts text from supported file types.

    Supported:
    - .txt
    - .md
    - .pdf
    - .docx
    """

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}

    def load(self, path: Path) -> list[DocumentPage]:
        path = Path(path)
        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")

        if extension in {".txt", ".md"}:
            return self._load_text_file(path)

        if extension == ".pdf":
            return self._load_pdf(path)

        if extension == ".docx":
            return self._load_docx(path)

        raise ValueError(f"Unsupported file type: {extension}")

    def _load_text_file(self, path: Path) -> list[DocumentPage]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return [DocumentPage(page_number=1, text=text)]

    def _load_pdf(self, path: Path) -> list[DocumentPage]:
        pages: list[DocumentPage] = []

        with fitz.open(path) as pdf:
            for index, page in enumerate(pdf, start=1):
                text = page.get_text("text")
                pages.append(DocumentPage(page_number=index, text=text))

        return pages

    def _load_docx(self, path: Path) -> list[DocumentPage]:
        document = docx.Document(path)

        parts: list[str] = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())

        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)

        return [DocumentPage(page_number=1, text="\n".join(parts))]